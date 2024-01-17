from project import get_path, merger, text_counter, Path
from docx import Document
from unittest.mock import patch
import pytest
import docx
import os

working_directory = os.getcwd()
     
@patch("builtins.input", lambda _: "merge!" )
@patch("project.Path")
def test_get_path(mock_class):
    mock_class_instance = mock_class.return_value 
    mock_class_instance.action.return_value = \
        "../test"
    mock_class_instance.return_key.return_value = "merge!"
    mock_class_instance.decision.return_value = ""
    assert get_path() == "../test"

    mock_class_instance = mock_class.return_value 
    mock_class_instance.action.return_value = "D:/ABC"
    mock_class_instance.return_key.return_value = "merge!"
    mock_class_instance.decision.return_value = ""
    with pytest.raises(AssertionError):
        assert get_path() == "D:/DEF"

def test_merger():
    assert type(merger(os.path.join(working_directory, "test"))) == str
    with pytest.raises(docx.opc.exceptions.PackageNotFoundError):
        assert merger(Document("dog"))

def test_text_counter():
    assert type(text_counter(Document(
        os.path.join(working_directory, "test/test1.docx"
    )))) == tuple
    with pytest.raises(docx.opc.exceptions.PackageNotFoundError):
        assert text_counter(Document("dog"))
