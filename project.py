"""DOCX.MERGER"""
""" This program merges multiple .docx files into one."""

# Importing modules needed for the job:
# "Python-docx" is the engine for reading and writing Word documents.
# "Docxcompose" merges Word files.
# "OS" does all the necessary operations with filepaths.
# "Datetime" adds today's date to the merged file's name.
# "Sys" allows the user to terminate the program.
import os
from datetime import date
import sys
from docx import Document
from docxcompose.composer import Composer

CRED = "\33[31m"
CGREEN = "\33[32m"
CYELLOW = "\33[33m"
CEND = "\33[0m"

class Path:
    """This class makes it easier to manage filepaths."""
    def __init__(self, key):
        self.key = key

    def return_key(self):
        """Returns the value of input, i. e., keyword."""
        return self.key

    def action(self, path, dirlist):
        """Lists keywords and respective actions when navigating through
        the file system."""
        if self.key == "cd!":
            return os.path.split(path)[0]
        elif self.key in dirlist:
            return os.path.join(path, self.key)
        elif self.key == "help!":
            print(f"{CYELLOW}-------------HELP---------------\n"
                  "\"cd!\" jumps one directory higher.\n"
                  "Typing in the directory name jumps into it.\n"
                  "\"merge!\" launches the program.\n"
                  f"\"exit!\" terminates the program.{CEND}")
            return path
        elif self.key == "merge!":
            return path
        elif self.key == "exit!":
            sys.exit(f"{CRED}Program terminated.{CEND}")
        else:
            print(f"{CRED}Not a directory.{CEND}")
            return path

    def decision(self, path):
        """Prints out the files to be merged
        and asks the user for confirmation."""
        filelist = [file for file in os.listdir(path)
                    if file.endswith(".docx")]
        if len(filelist) == 0:
            sys.exit("No files to merge.")
        print("Following files will be merged:\n"
              "-----------------------")
        for file in filelist:
            print(file)
        while True:
            conf = input("------------------------\n"
                         "Proceed? Press Enter or type \"exit!\".\n")
            if conf == "":
                break
            elif conf == "exit!":
                sys.exit(f"{CRED}Program terminated.{CEND}")
            else:
                pass


def main():
    """The batch for all the secondary functions. Gets the path to the selected
    directory, merges all the .docx files there, and counts its words, 
    characters and \"standard pages\"."""
    full_path = get_path()
    print(merger(full_path))

def get_path():
    """Allows user to locate the folder with the files to be merged."""
    absolute = os.path.dirname(__file__)
    while True:
        print("---------------------------------\n"
            f"Your current folder:\n{absolute}")
        dirlist = [f.name for f in os.scandir(absolute) if f.is_dir()]
        if len(dirlist) > 0:
            for folder in dirlist:
                print(f"+ {folder}", end = "\n")
        else:
            print(f"{CYELLOW}No subdirectories.{CEND}")
        key = Path(input("Navigate to your folder of .docx files. "
                              "(Use \"help!\" for more info):\n"))
        absolute = key.action(absolute, dirlist)
        if key.return_key() == "merge!":
            key.decision(absolute)
            break
    return absolute

def merger(full_path):
    """Merges all the .docx files in the directory into a new file."""
    folder = [file for file in os.listdir(full_path) if file.endswith(".docx")]
    master = Document(os.path.join(full_path, folder[0]))
    composer = Composer(master)
    for i in range(1, len(folder)):
        document = Document(os.path.join(full_path, folder[i]))
        composer.append(document)
    composer.save(
        os.path.join(full_path, f"merged_{date.today()}_{folder[0]}"
                     ))
    new_doc = Document(
        os.path.join(full_path, f"merged_{date.today()}_{folder[0]}"
                     ))
    words, chars, ns = text_counter(new_doc)
    return f"{CGREEN}--------------------------------\n" \
            f"All merged! Check the folder.\n" \
            f"Word count: {words}\n" \
            f"Character count: {chars}\n" \
            f"Number of NS (rounded up): {ns}{CEND}"

def text_counter(new_doc):
    """Counts words, characters, and standard pages on the output."""
    word_count = sum(len(p.text.split()) for p in new_doc.paragraphs)
    chars_count = sum(len(p.text) for p in new_doc.paragraphs)
    num_of_ns = round((chars_count / 1800) + (chars_count % 1800 > 0))
    return word_count, chars_count, num_of_ns

if __name__ == "__main__":
    main()
