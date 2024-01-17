<<<<<<< HEAD
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
import os
from datetime import date
import sys

class Path:
    def __init__(self, path):
        self.path = path

    def cd(self, key, path):
        if key == "cd!":
            return os.path.split(path)[0]



def main():  
    full_path = get_path()
    new_doc = get_new_doc()
    print(merger(full_path, new_doc))
    

    # Need TO CONSERVE FORMATTING (NOW DUPLICATING)
    # Need to ensure simplicity and user friendliness:
        # MAKE IT INTO CLASSES
    # Need to catch out the exceptions 

def get_path():
    absolute_path = os.path.dirname(__file__)
    while True:
        print(f"Your current folder:\n{absolute_path}")
        dirlist = [f.name for f in os.scandir(absolute_path) if f.is_dir()]
        if len(dirlist) > 0:
            for dir in dirlist:
                print(f"+ {dir}")
        else:
            print("No subdirectories.")
        relative_path = input("Navigate to your folder of .docx files. "
                              "(Use \"help!\" for more info):\n")
         
        if relative_path == "cd!":
            absolute_path = os.path.split(absolute_path)[0]
        elif relative_path in dirlist:
            absolute_path = os.path.join(absolute_path, relative_path)
        elif relative_path == "help!":
            print("---HELP---\n"
                  "\"cd!\" jumps one directory higher.\n"
                  "Typing in the directory name jumps into it.\n"
                  "\"merge!\" launches the program.\n"
                  "\"exit!\" terminates the program.\n"
                  "---")
        elif relative_path == "merge!":
            path = absolute_path
            break
        elif relative_path == "exit!":
            sys.exit("Program terminated.")
        else:
            print("Not a directory.")
            pass 
    return path

def get_new_doc():
    new_doc = Document()
    style = new_doc.styles['Normal']
    font = style.font
    font.name = "Times"
    font.size = Pt(12)
    sections = new_doc.sections
    for section in sections:
        margin = Cm(2.5)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
    return new_doc

def merger(full_path, new_doc):
    dir = os.listdir(full_path)
    for file in dir:
        try:
            file = Document(f"{full_path}\\{file}")
            for paragraph in file.paragraphs:
                new_doc.add_paragraph(paragraph.text)
                for run in paragraph.runs:
                    if "w:br" in run._element.xml \
                    and "w:type=\"page\"" in run._element.xml:
                        new_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                    if "w:highlight" in run._element.xml:
                        new_doc.add_paragraph().add_run(
                            run.text).font.highlight_color = WD_COLOR_INDEX.RED
                    if run.italic:
                        new_doc.add_paragraph().add_run(
                            run.text).font.italic = True
        except docx.opc.exceptions.PackageNotFoundError:
            pass

    words, chars = text_counter(new_doc)
    new_doc.save(f"{full_path}\\merged_{date.today()}_{dir[0]}.docx")
    return f"All merged!\nWord count: {words}\nCharacter count: {chars}"

def text_counter(new_doc):
    word_count = sum(len(p.text.split()) for p in new_doc.paragraphs)    
    chars_count = sum(len(p.text) for p in new_doc.paragraphs)
    return word_count, chars_count
              

if __name__ == "__main__":
    main()
=======
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
import os
from datetime import date
import sys

class Path:
    def __init__(self, path):
        self.path = path

    def cd(self, key, path):
        if key == "cd!":
            return os.path.split(path)[0]



def main():  
    full_path = get_path()
    new_doc = get_new_doc()
    print(merger(full_path, new_doc))
    

    # Need TO CONSERVE FORMATTING (NOW DUPLICATING)
    # Need to ensure simplicity and user friendliness:
        # MAKE IT INTO CLASSES
    # Need to catch out the exceptions 

def get_path():
    absolute_path = os.path.dirname(__file__)
    while True:
        print(f"Your current folder:\n{absolute_path}")
        dirlist = [f.name for f in os.scandir(absolute_path) if f.is_dir()]
        if len(dirlist) > 0:
            for dir in dirlist:
                print(f"+ {dir}")
        else:
            print("No subdirectories.")
        relative_path = input("Navigate to your folder of .docx files. "
                              "(Use \"help!\" for more info):\n")
         
        if relative_path == "cd!":
            absolute_path = os.path.split(absolute_path)[0]
        elif relative_path in dirlist:
            absolute_path = os.path.join(absolute_path, relative_path)
        elif relative_path == "help!":
            print("---HELP---\n"
                  "\"cd!\" jumps one directory higher.\n"
                  "Typing in the directory name jumps into it.\n"
                  "\"merge!\" launches the program.\n"
                  "\"exit!\" terminates the program.\n"
                  "---")
        elif relative_path == "merge!":
            path = absolute_path
            break
        elif relative_path == "exit!":
            sys.exit("Program terminated.")
        else:
            print("Not a directory.")
            pass 
    return path

def get_new_doc():
    new_doc = Document()
    style = new_doc.styles['Normal']
    font = style.font
    font.name = "Times"
    font.size = Pt(12)
    sections = new_doc.sections
    for section in sections:
        margin = Cm(2.5)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
    return new_doc

def merger(full_path, new_doc):
    dir = os.listdir(full_path)
    for file in dir:
        try:
            file = Document(f"{full_path}\\{file}")
            for paragraph in file.paragraphs:
                new_doc.add_paragraph(paragraph.text)
                for run in paragraph.runs:
                    if "w:br" in run._element.xml \
                    and "w:type=\"page\"" in run._element.xml:
                        new_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                    if "w:highlight" in run._element.xml:
                        new_doc.add_paragraph().add_run(
                            run.text).font.highlight_color = WD_COLOR_INDEX.RED
                    if run.italic:
                        new_doc.add_paragraph().add_run(
                            run.text).font.italic = True
        except docx.opc.exceptions.PackageNotFoundError:
            pass

    words, chars = text_counter(new_doc)
    new_doc.save(f"{full_path}\\merged_{date.today()}_{dir[0]}.docx")
    return f"All merged!\nWord count: {words}\nCharacter count: {chars}"

def text_counter(new_doc):
    word_count = sum(len(p.text.split()) for p in new_doc.paragraphs)    
    chars_count = sum(len(p.text) for p in new_doc.paragraphs)
    return word_count, chars_count
              

if __name__ == "__main__":
    main()
>>>>>>> da2dcdd2c1dc555ffc5b8e14dd2d2b422a61ee05
