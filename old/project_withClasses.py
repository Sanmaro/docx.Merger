<<<<<<< HEAD
""" This program is intended to merge multiple .docx files into one."""


# TO DO LIST:
    # Need TO CONSERVE FORMATTING (NOW DUPLICATING)
    # Need to ensure simplicity and user friendliness:
        # MAKE IT INTO CLASSES
    # Need to catch out the exceptions





# Importing modules needed for the job.
# "Docx" being the engine for reading and writing Word documents.
# "OS" doing all necessary operations with filepaths.
# "Datetime" adding today's date to the merged file's name.
# "Sys" allowing the user to terminate the program.
import os
from datetime import date
import sys
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX

class Path:
    """This classis making it easier to manage filepaths."""
    def __init__(self, key):
        self.key = key

    def return_key(self):
        """Returns the value of input, i. e., keyword."""
        return self.key

    def action(self, path, dirlist):
        """List of keywords and respective actions when navigating through
        the file system."""
        if self.key == "cd!":
            return os.path.split(path)[0]
        elif self.key in dirlist:
            return os.path.join(path, self.key)
        elif self.key == "help!":
            print("-------------HELP---------------\n"
                  "\"cd!\" jumps one directory higher.\n"
                  "Typing in the directory name jumps into it.\n"
                  "\"merge!\" launches the program.\n"
                  "\"exit!\" terminates the program.\n"
                  "----------------------------")
            return path
        elif self.key == "merge!":
            return path
        elif self.key == "exit!":
            sys.exit("Program terminated.")
        else:
            print("Not a directory.")
            return path

def main():
    """The batch for all the secondary functions. Gets the path to the selected
    directory, creates a new document and merges all the .docx files 
    from the folder into it."""
    full_path = get_path()
    new_doc = get_new_doc()
    print(merger(full_path, new_doc))

def get_path():
    """Allows user to find the folder with the files to be merged."""
    absolute = os.path.dirname(__file__)
    while True:
        print(f"Your current folder:\n{absolute}")
        dirlist = [f.name for f in os.scandir(absolute) if f.is_dir()]
        if len(dirlist) > 0:
            for folder in dirlist:
                print(f"+ {folder}", end = "\n")
        else:
            print("No subdirectories.")
        key = Path(input("Navigate to your folder of .docx files. "
                              "(Use \"help!\" for more info):\n"))        
        absolute = key.action(absolute, dirlist)
        if key.return_key() == "merge!":
            break
    return absolute

def get_new_doc():
    """Creates a new, blank Word document."""
    new_doc = Document()
    style = new_doc.styles['Normal']
    font = style.font
    font.name = "Times"
    font.size = Pt(12)
    sections = new_doc.sections
    for section in sections:
        margin = Cm(2.5)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
    return new_doc

def merger(full_path, new_doc):
    """Merges all the .docx files in the directory, 
    using the newly created document."""
    folder = os.listdir(full_path)
    for file in folder:
        try:
            file = Document(f"{full_path}\\{file}")
            for paragraph in file.paragraphs:
                # new_doc.add_paragraph(paragraph.text)
                for run in paragraph.runs:
                    if "w:br" in run._element.xml \
                    and "w:type=\"page\"" in run._element.xml:
                        new_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                    elif "w:highlight" in run._element.xml:
                        new_doc.add_paragraph().add_run(
                            run.text).font.highlight_color = WD_COLOR_INDEX.RED
                    elif run.italic:
                        new_doc.add_paragraph().add_run(
                            run.text).font.italic = True
                    else:
                        new_doc.add_paragraph().add_run(run.text)
        except docx.opc.exceptions.PackageNotFoundError:
            pass
    words, chars = text_counter(new_doc)
    new_doc.save(f"{full_path}\\merged_{date.today()}_{folder[0]}")
    return f"All merged!\nWord count: {words}\nCharacter count: {chars}"

def text_counter(new_doc):
    """Counts words and characters on the output"""
    word_count = sum(len(p.text.split()) for p in new_doc.paragraphs)
    chars_count = sum(len(p.text) for p in new_doc.paragraphs)
    return word_count, chars_count

if __name__ == "__main__":
    main()
=======
""" This program is intended to merge multiple .docx files into one."""


# TO DO LIST:
    # Need TO CONSERVE FORMATTING (NOW DUPLICATING)
    # Need to ensure simplicity and user friendliness:
        # MAKE IT INTO CLASSES
    # Need to catch out the exceptions





# Importing modules needed for the job.
# "Docx" being the engine for reading and writing Word documents.
# "OS" doing all necessary operations with filepaths.
# "Datetime" adding today's date to the merged file's name.
# "Sys" allowing the user to terminate the program.
import os
from datetime import date
import sys
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX

class Path:
    """This classis making it easier to manage filepaths."""
    def __init__(self, key):
        self.key = key

    def return_key(self):
        """Returns the value of input, i. e., keyword."""
        return self.key

    def action(self, path, dirlist):
        """List of keywords and respective actions when navigating through
        the file system."""
        if self.key == "cd!":
            return os.path.split(path)[0]
        elif self.key in dirlist:
            return os.path.join(path, self.key)
        elif self.key == "help!":
            print("-------------HELP---------------\n"
                  "\"cd!\" jumps one directory higher.\n"
                  "Typing in the directory name jumps into it.\n"
                  "\"merge!\" launches the program.\n"
                  "\"exit!\" terminates the program.\n"
                  "----------------------------")
            return path
        elif self.key == "merge!":
            return path
        elif self.key == "exit!":
            sys.exit("Program terminated.")
        else:
            print("Not a directory.")
            return path

def main():
    """The batch for all the secondary functions. Gets the path to the selected
    directory, creates a new document and merges all the .docx files 
    from the folder into it."""
    full_path = get_path()
    new_doc = get_new_doc()
    print(merger(full_path, new_doc))

def get_path():
    """Allows user to find the folder with the files to be merged."""
    absolute = os.path.dirname(__file__)
    while True:
        print(f"Your current folder:\n{absolute}")
        dirlist = [f.name for f in os.scandir(absolute) if f.is_dir()]
        if len(dirlist) > 0:
            for folder in dirlist:
                print(f"+ {folder}", end = "\n")
        else:
            print("No subdirectories.")
        key = Path(input("Navigate to your folder of .docx files. "
                              "(Use \"help!\" for more info):\n"))        
        absolute = key.action(absolute, dirlist)
        if key.return_key() == "merge!":
            break
    return absolute

def get_new_doc():
    """Creates a new, blank Word document."""
    new_doc = Document()
    style = new_doc.styles['Normal']
    font = style.font
    font.name = "Times"
    font.size = Pt(12)
    sections = new_doc.sections
    for section in sections:
        margin = Cm(2.5)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
    return new_doc

def merger(full_path, new_doc):
    """Merges all the .docx files in the directory, 
    using the newly created document."""
    folder = os.listdir(full_path)
    for file in folder:
        try:
            file = Document(f"{full_path}\\{file}")
            for paragraph in file.paragraphs:
                # new_doc.add_paragraph(paragraph.text)
                for run in paragraph.runs:
                    if "w:br" in run._element.xml \
                    and "w:type=\"page\"" in run._element.xml:
                        new_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                    elif "w:highlight" in run._element.xml:
                        new_doc.add_paragraph().add_run(
                            run.text).font.highlight_color = WD_COLOR_INDEX.RED
                    elif run.italic:
                        new_doc.add_paragraph().add_run(
                            run.text).font.italic = True
                    else:
                        new_doc.add_paragraph().add_run(run.text)
        except docx.opc.exceptions.PackageNotFoundError:
            pass
    words, chars = text_counter(new_doc)
    new_doc.save(f"{full_path}\\merged_{date.today()}_{folder[0]}")
    return f"All merged!\nWord count: {words}\nCharacter count: {chars}"

def text_counter(new_doc):
    """Counts words and characters on the output"""
    word_count = sum(len(p.text.split()) for p in new_doc.paragraphs)
    chars_count = sum(len(p.text) for p in new_doc.paragraphs)
    return word_count, chars_count

if __name__ == "__main__":
    main()
>>>>>>> da2dcdd2c1dc555ffc5b8e14dd2d2b422a61ee05
